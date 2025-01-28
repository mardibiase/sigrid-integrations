# Copyright Software Improvement Group
#
# Licensed under the Apache License, Version 2.0 (the "License");
# you may not use this file except in compliance with the License.
# You may obtain a copy of the License at
# 
#     http://www.apache.org/licenses/LICENSE-2.0
#
# Unless required by applicable law or agreed to in writing, software
# distributed under the License is distributed on an "AS IS" BASIS,
# WITHOUT WARRANTIES OR CONDITIONS OF ANY KIND, either express or implied.
# See the License for the specific language governing permissions and
# limitations under the License.

from .report import Report

from .word_utils import WordUtils
from .osh_report import OSHReport, OSHData
import docx, re, logging

class WordReport(Report):

    def __init__(self, customer, system, token, out_file, template):
        super().__init__(customer, system, token)
        self.out_file = out_file
        self.document = docx.Document(template)

    def save_file(self):
        out = self.out_file.replace("pptx", "docx")
        logging.info(f"Written output to {out}")
        self.document.save(out)

    def create_architecture_quality_report_specific(self, architecture_data):
        pass

    def create_maintainability_report_specific(self, maint_data):
        pass

    def create_osh_report_specific(self, osh_data: OSHData, osh_report: OSHReport):
        pass

    def create_technology_report_specific(self, unsorted_tech_data, sorted_tech_data):
        pass

    def fill_metric_report_specific(self, model, metric, value):
        pass

    def create_architecture_quality_report_specific(self, architecture_metrics):
        pass

    def update_placeholder(self, placeholder, replacement_text):
        paragraphs = self.find_in_document(placeholder)
        for paragraph in paragraphs:
            WordUtils.merge_runs_with_same_formatting(paragraph)
            self.update_paragraph(paragraph, placeholder, replacement_text)
            #run.text = run.text.replace(placeholder, str(value))

    def update_paragraph(self, paragraph, placeholder, replacement_text):

        run_with_placeholder = None
        if paragraph.runs:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run_with_placeholder = run
                    break
        
        if run_with_placeholder:
            run_with_placeholder.text = run_with_placeholder.text.replace(placeholder, str(replacement_text))
            logging.debug(f"Replacing: {placeholder} with \"{replacement_text}\". New text: {run_with_placeholder.text}")
        
    def find_in_document(self, placeholder):
        paragraphs = []
        
        for paragraph in self.document.paragraphs:
            for run in paragraph.runs:
                if re.match(rf".*\b{placeholder}\b.*", run.text):
                    paragraphs.append(paragraph)

        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if re.match(rf".*\b{placeholder}\b.*", run.text):
                                paragraphs.append(paragraph)

        return paragraphs